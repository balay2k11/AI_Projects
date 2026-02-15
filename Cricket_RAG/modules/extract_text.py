from __future__ import annotations
from pathlib import Path
import pdfplumber
from docx import Document

def extract_text_from_pdf(path: Path) -> str:
    chunks = []
    with pdfplumber.open(str(path)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text() or ""
            if txt.strip():
                chunks.append(txt)
    return "\n".join(chunks).strip()

def extract_text_from_docx(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts).strip()

def extract_text(file_path: Path) -> str:
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    if ext == ".docx":
        return extract_text_from_docx(file_path)
    raise ValueError(f"Unsupported file type: {ext}")
